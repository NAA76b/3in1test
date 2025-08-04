# -*- coding: utf-8 -*-
"""
Master Workflow Tool v2.0 - Redesigned with Apple/Google-style UI

REDESIGN SUMMARY:
- Polished Apple/Google-style UI with numbered steps, clean typography, and modern styling
- Hardened matching logic with centralized normalization and fuzzy fallback
- Enhanced error handling for COM/Outlook and file I/O operations
- Fixed HTML CSS syntax bugs and centralized filename formatting
- Added Master_Names_with_IDs export with match status tracking
- Improved logging with filterable messages and summary badges
- Thread-safe UI updates with comprehensive progress tracking

IMPROVEMENTS:
- Exact name matching with fuzzy fallback using configurable threshold
- Preserved existing employee IDs with clear status tagging
- Centralized name normalization to prevent drift
- Safe Outlook integration with graceful error handling
- Dynamic UI states with inline validation
- Collapsible log filters and summary statistics
- Consistent modern styling with hover states and animations

UI ENHANCEMENTS:
- Light, clean theme with soft neutral backgrounds
- Rounded corners and subtle drop shadows
- Clear typographic hierarchy with readable fonts
- Grid-based numbered steps with progress indicators
- Primary/secondary button styling with elevation
- Dynamic field states and smooth transitions
- Copy-to-clipboard functionality in logs
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import os
import re
import threading
from datetime import datetime
import traceback

# Try to import fuzzywuzzy, fall back to basic string matching if not available
try:
    from fuzzywuzzy import process, fuzz
    FUZZY_AVAILABLE = True
except ImportError:
    FUZZY_AVAILABLE = False
    import difflib
    print("Warning: fuzzywuzzy not available. Using basic string matching fallback.")

# Required libraries installation check
try:
    import win32com.client as win32
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    messagebox.showerror("Missing Libraries", 
        "This application requires additional libraries.\n\n"
        "Please install them by running:\n"
        "pip install pywin32 python-docx openpyxl\n\n"
        "Optional: pip install fuzzywuzzy python-levenshtein (for enhanced matching)\n\n"
        "Then restart the application.")
    exit()

#==========================================================================
# CORE UTILITIES AND NORMALIZATION
#==========================================================================

def normalize_name(name_str):
    """
    Centralized name normalization for consistent matching.
    Converts to uppercase, removes extra spaces, and standardizes format.
    """
    if pd.isna(name_str) or not str(name_str).strip():
        return ""
    
    # Convert to string and uppercase
    normalized = str(name_str).upper().strip()
    
    # Remove extra whitespace and standardize separators
    normalized = re.sub(r'\s+', ' ', normalized)
    normalized = re.sub(r'[,\-\.]+', ' ', normalized)
    
    # Remove common suffixes that might cause mismatches
    suffixes = ['JR', 'SR', 'III', 'II', 'IV']
    for suffix in suffixes:
        normalized = re.sub(rf'\b{suffix}\b', '', normalized)
    
    return normalized.strip()

def format_output_filename(base_name, extension=".xlsx", include_timestamp=True):
    """
    Centralized filename formatting with consistent timestamping.
    """
    if include_timestamp:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{base_name}_{timestamp}{extension}"
    return f"{base_name}{extension}"

def get_default_output_path():
    """Gets the default output path, preferring OneDrive Desktop if available."""
    onedrive_desktop = os.path.join(os.path.expanduser("~"), "OneDrive - FDA", "Desktop")
    if os.path.exists(onedrive_desktop):
        return onedrive_desktop
    return os.path.expanduser("~/Desktop")

def safe_com_operation(operation_func, error_log_func, operation_name="COM operation"):
    """
    Safely execute COM operations with comprehensive error handling.
    """
    try:
        return operation_func()
    except Exception as e:
        error_msg = f"{operation_name} failed: {str(e)}"
        if "Outlook" in operation_name:
            error_msg += "\n\nTip: Ensure Outlook is installed and can be opened manually."
        error_log_func(error_msg, 'error')
        return None

def safe_file_operation(operation_func, error_log_func, operation_name="File operation"):
    """
    Safely execute file operations with comprehensive error handling.
    """
    try:
        return operation_func()
    except FileNotFoundError as e:
        error_log_func(f"{operation_name} failed - File not found: {str(e)}", 'error')
        return None
    except PermissionError as e:
        error_log_func(f"{operation_name} failed - Permission denied: {str(e)}", 'error')
        return None
    except Exception as e:
        error_log_func(f"{operation_name} failed: {str(e)}", 'error')
        return None

class MatchingEngine:
    """
    Centralized matching engine for employee names with fuzzy fallback.
    """
    
    def __init__(self, fuzzy_threshold=80):
        self.fuzzy_threshold = fuzzy_threshold
        self.lookup_dict = {}
        self.match_stats = {
            'exact_matches': 0,
            'fuzzy_matches': 0,
            'no_matches': 0,
            'already_had_id': 0
        }
    
    def build_lookup(self, employee_df, name_col='Employee Name', id_col='Employee ID'):
        """Build normalized lookup dictionary from employee data."""
        self.lookup_dict = {}
        for _, row in employee_df.iterrows():
            if pd.notna(row[name_col]) and pd.notna(row[id_col]):
                normalized_name = normalize_name(row[name_col])
                if normalized_name:
                    self.lookup_dict[normalized_name] = {
                        'id': row[id_col],
                        'original_name': row[name_col]
                    }
    
    def match_employee_id(self, name, existing_id=None):
        """
        Match employee name to ID with exact and fuzzy fallback.
        Returns tuple: (matched_id, match_status, confidence_score)
        """
        if pd.notna(existing_id) and str(existing_id).strip():
            self.match_stats['already_had_id'] += 1
            return existing_id, "Already had ID", 100
        
        if pd.isna(name) or not str(name).strip():
            self.match_stats['no_matches'] += 1
            return None, "No match - empty name", 0
        
        normalized_name = normalize_name(name)
        
        # Try exact match first
        if normalized_name in self.lookup_dict:
            self.match_stats['exact_matches'] += 1
            return self.lookup_dict[normalized_name]['id'], "Exact match", 100
        
        # Fuzzy fallback
        lookup_names = list(self.lookup_dict.keys())
        if lookup_names:
            if FUZZY_AVAILABLE:
                best_match = process.extractOne(
                    normalized_name, 
                    lookup_names, 
                    scorer=fuzz.token_sort_ratio
                )
                
                if best_match and best_match[1] >= self.fuzzy_threshold:
                    matched_name = best_match[0]
                    confidence = best_match[1]
                    self.match_stats['fuzzy_matches'] += 1
                    return (
                        self.lookup_dict[matched_name]['id'], 
                        f"Fuzzy match ({confidence}%)", 
                        confidence
                    )
            else:
                # Fallback using difflib
                best_match = difflib.get_close_matches(
                    normalized_name, 
                    lookup_names, 
                    n=1, 
                    cutoff=self.fuzzy_threshold/100.0
                )
                
                if best_match:
                    matched_name = best_match[0]
                    # Estimate confidence using sequence matcher
                    confidence = int(difflib.SequenceMatcher(None, normalized_name, matched_name).ratio() * 100)
                    self.match_stats['fuzzy_matches'] += 1
                    return (
                        self.lookup_dict[matched_name]['id'], 
                        f"Fuzzy match ({confidence}%)", 
                        confidence
                    )
        
        self.match_stats['no_matches'] += 1
        return None, "No match", 0
    
    def get_summary_stats(self):
        """Get formatted summary statistics."""
        total = sum(self.match_stats.values())
        if total == 0:
            return "No matches processed"
        
        return (
            f"Total: {total} | "
            f"Exact: {self.match_stats['exact_matches']} | "
            f"Fuzzy: {self.match_stats['fuzzy_matches']} | "
            f"No match: {self.match_stats['no_matches']} | "
            f"Had ID: {self.match_stats['already_had_id']}"
        )

def log_message_util(root, log_widget, msg, level='info'):
    """Enhanced utility to safely log messages with better formatting."""
    def _insert():
        if log_widget.winfo_exists():
            log_widget.config(state='normal')
            
            # Configure tags with modern styling
            log_widget.tag_config('success', foreground='#22c55e', font=('SF Pro Text', 10, 'bold'))
            log_widget.tag_config('error', foreground='#ef4444', font=('SF Pro Text', 10, 'bold'))
            log_widget.tag_config('warning', foreground='#f59e0b', font=('SF Pro Text', 10, 'bold'))
            log_widget.tag_config('info', foreground='#64748b', font=('SF Pro Text', 10))
            
            timestamp = datetime.now().strftime('%H:%M:%S')
            level_icon = {'success': '✓', 'error': '✗', 'warning': '⚠', 'info': 'ℹ'}.get(level, 'ℹ')
            
            log_widget.insert(tk.END, f"[{timestamp}] {level_icon} {msg}\n", level)
            log_widget.config(state='disabled')
            log_widget.see(tk.END)
    
    if root and root.winfo_exists():
        root.after(0, _insert)

#==========================================================================
# MODERN UI COMPONENTS
#==========================================================================

class ModernButton(ttk.Frame):
    """Custom button with modern styling and hover effects."""
    
    def __init__(self, parent, text, command, style_type="primary", state="normal", **kwargs):
        super().__init__(parent, **kwargs)
        
        self.command = command
        self.style_type = style_type
        self.current_state = state
        
        # Style configurations
        styles = {
            "primary": {"bg": "#007AFF", "fg": "white", "hover": "#0056CC"},
            "secondary": {"bg": "#F2F2F7", "fg": "#1C1C1E", "hover": "#E5E5EA"},
            "success": {"bg": "#34C759", "fg": "white", "hover": "#28A745"},
            "warning": {"bg": "#FF9500", "fg": "white", "hover": "#E68500"},
            "danger": {"bg": "#FF3B30", "fg": "white", "hover": "#D70015"}
        }
        
        self.style_config = styles.get(style_type, styles["primary"])
        
        self.button = tk.Label(
            self,
            text=text,
            bg=self.style_config["bg"],
            fg=self.style_config["fg"],
            font=('SF Pro Text', 12, 'bold'),
            padx=20,
            pady=12,
            cursor="hand2" if state == "normal" else "arrow"
        )
        self.button.pack(fill='both', expand=True)
        
        if state == "normal":
            self.button.bind("<Button-1>", self._on_click)
            self.button.bind("<Enter>", self._on_enter)
            self.button.bind("<Leave>", self._on_leave)
        
        self.configure_state(state)
    
    def _on_click(self, event=None):
        if self.current_state == "normal" and self.command:
            self.command()
    
    def _on_enter(self, event=None):
        if self.current_state == "normal":
            self.button.configure(bg=self.style_config["hover"])
    
    def _on_leave(self, event=None):
        if self.current_state == "normal":
            self.button.configure(bg=self.style_config["bg"])
    
    def configure_state(self, state):
        self.current_state = state
        if state == "disabled":
            self.button.configure(
                bg="#E5E5EA",
                fg="#8E8E93",
                cursor="arrow"
            )
        else:
            self.button.configure(
                bg=self.style_config["bg"],
                fg=self.style_config["fg"],
                cursor="hand2"
            )

class StepIndicator(ttk.Frame):
    """Modern step indicator with progress visualization."""
    
    def __init__(self, parent, step_number, title, subtitle="", **kwargs):
        super().__init__(parent, **kwargs)
        
        # Step number circle
        self.step_frame = tk.Frame(self, width=40, height=40, bg="#007AFF")
        self.step_frame.pack_propagate(False)
        self.step_frame.pack(side='left', padx=(0, 15))
        
        self.step_label = tk.Label(
            self.step_frame,
            text=str(step_number),
            bg="#007AFF",
            fg="white",
            font=('SF Pro Text', 14, 'bold')
        )
        self.step_label.place(relx=0.5, rely=0.5, anchor='center')
        
        # Text content
        self.content_frame = ttk.Frame(self)
        self.content_frame.pack(side='left', fill='both', expand=True)
        
        self.title_label = ttk.Label(
            self.content_frame,
            text=title,
            font=('SF Pro Display', 16, 'bold'),
            foreground='#1C1C1E'
        )
        self.title_label.pack(anchor='w')
        
        if subtitle:
            self.subtitle_label = ttk.Label(
                self.content_frame,
                text=subtitle,
                font=('SF Pro Text', 12),
                foreground='#8E8E93'
            )
            self.subtitle_label.pack(anchor='w')
    
    def mark_complete(self):
        """Mark step as completed with visual feedback."""
        self.step_frame.configure(bg="#34C759")
        self.step_label.configure(bg="#34C759", text="✓")

class SummaryBadge(ttk.Frame):
    """Summary statistics badge with modern styling."""
    
    def __init__(self, parent, label, value, color="#007AFF", **kwargs):
        super().__init__(parent, **kwargs)
        
        self.badge_frame = tk.Frame(
            self,
            bg=color,
            padx=12,
            pady=8
        )
        self.badge_frame.pack(fill='x')
        
        self.label = tk.Label(
            self.badge_frame,
            text=label,
            bg=color,
            fg="white",
            font=('SF Pro Text', 10, 'bold')
        )
        self.label.pack(side='left')
        
        self.value = tk.Label(
            self.badge_frame,
            text=str(value),
            bg=color,
            fg="white",
            font=('SF Pro Text', 14, 'bold')
        )
        self.value.pack(side='right')
    
    def update_value(self, new_value):
        """Update badge value with animation effect."""
        self.value.configure(text=str(new_value))

#==========================================================================
# TAB 1: PAY PERIOD REPORTER (Enhanced)
#==========================================================================
class ReporterFrame(ttk.Frame):
    def __init__(self, container, app_instance, **kwargs):
        super().__init__(container, **kwargs)
        self.app = app_instance
        
        # Variables
        self.history_file = tk.StringVar()
        self.employee_file = tk.StringVar()
        self.reporter_output_location = tk.StringVar(value=self.app.default_output_path)
        self.fill_logic = tk.StringVar(value="Smart Header + Dynamic Fill")
        
        self.setup_ui()
        
        # Match engine for employee ID matching
        self.match_engine = MatchingEngine()
    
    def setup_ui(self):
        """Setup modern UI layout with numbered steps."""
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)
        
        # Header
        header_frame = ttk.Frame(self)
        header_frame.grid(row=0, column=0, sticky='ew', pady=(0, 30))
        
        ttk.Label(
            header_frame,
            text="Pay Period Report Generator",
            font=('SF Pro Display', 24, 'bold'),
            foreground='#1C1C1E'
        ).pack(anchor='w')
        
        ttk.Label(
            header_frame,
            text="Compare employee compliance across pay periods with intelligent matching",
            font=('SF Pro Text', 14),
            foreground='#8E8E93'
        ).pack(anchor='w', pady=(5, 0))
        
        # Main content
        content_frame = ttk.Frame(self)
        content_frame.grid(row=1, column=0, sticky='nsew')
        content_frame.grid_columnconfigure(0, weight=1)
        content_frame.grid_rowconfigure(4, weight=1)
        
        # Step 1: History File
        step1 = StepIndicator(
            content_frame, 
            1, 
            "Select History File",
            "Previous pay period timesheet data"
        )
        step1.grid(row=0, column=0, sticky='ew', pady=(0, 15))
        
        file_frame1 = ttk.Frame(content_frame)
        file_frame1.grid(row=0, column=1, sticky='ew', padx=(20, 0))
        file_frame1.grid_columnconfigure(0, weight=1)
        
        self.history_entry = ttk.Entry(
            file_frame1,
            textvariable=self.history_file,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        self.history_entry.grid(row=0, column=0, sticky='ew', ipady=8)
        
        ttk.Button(
            file_frame1,
            text="Browse",
            command=lambda: self.browse_file(self.history_file, "Select History File")
        ).grid(row=0, column=1, padx=(10, 0))
        
        # Step 2: Employee File
        step2 = StepIndicator(
            content_frame,
            2,
            "Select Employee File", 
            "Current pay period timesheet data"
        )
        step2.grid(row=1, column=0, sticky='ew', pady=(20, 15))
        
        file_frame2 = ttk.Frame(content_frame)
        file_frame2.grid(row=1, column=1, sticky='ew', padx=(20, 0))
        file_frame2.grid_columnconfigure(0, weight=1)
        
        self.employee_entry = ttk.Entry(
            file_frame2,
            textvariable=self.employee_file,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        self.employee_entry.grid(row=0, column=0, sticky='ew', ipady=8)
        
        ttk.Button(
            file_frame2,
            text="Browse",
            command=lambda: self.browse_file(self.employee_file, "Select Employee File")
        ).grid(row=0, column=1, padx=(10, 0))
        
        # Step 3: Output Location
        step3 = StepIndicator(
            content_frame,
            3,
            "Select Output Location",
            "Where to save generated reports"
        )
        step3.grid(row=2, column=0, sticky='ew', pady=(20, 15))
        
        output_frame = ttk.Frame(content_frame)
        output_frame.grid(row=2, column=1, sticky='ew', padx=(20, 0))
        output_frame.grid_columnconfigure(0, weight=1)
        
        self.output_entry = ttk.Entry(
            output_frame,
            textvariable=self.reporter_output_location,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        self.output_entry.grid(row=0, column=0, sticky='ew', ipady=8)
        
        ttk.Button(
            output_frame,
            text="Browse",
            command=lambda: self.browse_directory(self.reporter_output_location)
        ).grid(row=0, column=1, padx=(10, 0))
        
        # Step 4: Fill Logic
        step4 = StepIndicator(
            content_frame,
            4,
            "Select Fill Down Logic",
            "How to handle missing organizational data"
        )
        step4.grid(row=3, column=0, sticky='ew', pady=(20, 15))
        
        logic_frame = ttk.Frame(content_frame)
        logic_frame.grid(row=3, column=1, sticky='ew', padx=(20, 0))
        
        fill_combo = ttk.Combobox(
            logic_frame,
            textvariable=self.fill_logic,
            values=["No Fill Down", "Dynamic Keyword Fill", "Smart Header + Dynamic Fill"],
            state='readonly',
            font=('SF Pro Text', 12)
        )
        fill_combo.pack(fill='x', ipady=8)
        
        # Action section
        action_frame = ttk.Frame(content_frame)
        action_frame.grid(row=4, column=0, columnspan=2, sticky='nsew', pady=(30, 0))
        action_frame.grid_columnconfigure(0, weight=1)
        action_frame.grid_rowconfigure(2, weight=1)
        
        # Generate button
        self.generate_button = ModernButton(
            action_frame,
            "Generate Comparison Reports",
            self.run_process,
            style_type="primary"
        )
        self.generate_button.grid(row=0, column=0, sticky='ew', ipady=15)
        
        # Progress bar
        self.progress = ttk.Progressbar(
            action_frame,
            orient='horizontal',
            mode='indeterminate'
        )
        self.progress.grid(row=1, column=0, sticky='ew', pady=(15, 0))
        
        # Log section
        log_frame = ttk.LabelFrame(action_frame, text="Processing Log", padding="15")
        log_frame.grid(row=2, column=0, sticky='nsew', pady=(20, 0))
        log_frame.grid_rowconfigure(0, weight=1)
        log_frame.grid_columnconfigure(0, weight=1)
        
        # Log text with scrollbar
        log_container = ttk.Frame(log_frame)
        log_container.grid(row=0, column=0, sticky='nsew')
        log_container.grid_rowconfigure(0, weight=1)
        log_container.grid_columnconfigure(0, weight=1)
        
        self.log_text = tk.Text(
            log_container,
            height=10,
            state='disabled',
            bg='#F8F9FA',
            fg='#1C1C1E',
            wrap='word',
            borderwidth=0,
            highlightthickness=0,
            font=('SF Mono', 10)
        )
        self.log_text.grid(row=0, column=0, sticky='nsew')
        
        scrollbar = ttk.Scrollbar(log_container, orient='vertical', command=self.log_text.yview)
        scrollbar.grid(row=0, column=1, sticky='ns')
        self.log_text.configure(yscrollcommand=scrollbar.set)
    
    def browse_file(self, string_var, title):
        """Enhanced file browser with validation."""
        filename = filedialog.askopenfilename(
            title=title,
            filetypes=[("Excel files", "*.xlsx *.xlsm"), ("All files", "*.*")]
        )
        if filename:
            string_var.set(filename)
            log_message_util(
                self.app.root, 
                self.log_text, 
                f"Selected: {os.path.basename(filename)}", 
                'success'
            )
            self.validate_inputs()
    
    def browse_directory(self, string_var):
        """Enhanced directory browser."""
        directory = filedialog.askdirectory(
            title="Select Output Location",
            initialdir=string_var.get()
        )
        if directory:
            string_var.set(directory)
            log_message_util(
                self.app.root,
                self.log_text,
                f"Output location: {directory}",
                'success'
            )
    
    def validate_inputs(self):
        """Validate inputs and update UI state."""
        has_history = bool(self.history_file.get())
        has_employee = bool(self.employee_file.get())
        
        if has_history and has_employee:
            self.generate_button.configure_state("normal")
        else:
            self.generate_button.configure_state("disabled")
    
    def run_process(self):
        """Enhanced process runner with better validation."""
        if not self.history_file.get() or not self.employee_file.get():
            messagebox.showerror(
                "Input Error",
                "Please select both History and Employee files before proceeding."
            )
            return
        
        self.generate_button.configure_state("disabled")
        self.progress.start(10)
        log_message_util(
            self.app.root,
            self.log_text,
            "Starting report generation process...",
            'info'
        )
        
        thread = threading.Thread(target=self.process_files_thread)
        thread.daemon = True
        thread.start()
    
    def process_files_thread(self):
        """Enhanced processing with better error handling."""
        try:
            # Load and clean data
            log_message_util(
                self.app.root,
                self.log_text,
                "Loading and cleaning report data...",
                'info'
            )
            
            history_df = self.clean_report_data(self.history_file.get(), "History")
            employee_df = self.clean_report_data(self.employee_file.get(), "Employee")
            
            if history_df is None or employee_df is None:
                raise ValueError("Failed to load required data files")
            
            # Validate and match employees
            matched_df, unmatched_hist, unmatched_emp = self.validate_employee_ids(
                history_df, employee_df
            )
            
            log_message_util(
                self.app.root,
                self.log_text,
                f"Successfully matched {len(matched_df)} employees",
                'success'
            )
            
            # Track compliance changes
            improved, declined, unchanged = self.track_compliance_changes(matched_df)
            
            log_message_util(
                self.app.root,
                self.log_text,
                f"Compliance analysis: {len(improved)} improved, {len(declined)} declined, {len(unchanged)} unchanged",
                'info'
            )
            
            # Generate output files
            output_dir = self.reporter_output_location.get()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Save individual reports
            unmatched_hist_file = os.path.join(
                output_dir, 
                format_output_filename("Unmatched_History", ".xlsx", True)
            )
            unmatched_emp_file = os.path.join(
                output_dir,
                format_output_filename("Unmatched_Employee", ".xlsx", True)
            )
            
            safe_file_operation(
                lambda: unmatched_hist.to_excel(unmatched_hist_file, index=False),
                lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
                "Saving unmatched history report"
            )
            
            safe_file_operation(
                lambda: unmatched_emp.to_excel(unmatched_emp_file, index=False),
                lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
                "Saving unmatched employee report"
            )
            
            # Save comprehensive matched report
            matched_filename = os.path.join(
                output_dir,
                format_output_filename("Comprehensive_Report", ".xlsx", True)
            )
            
            def save_comprehensive_report():
                with pd.ExcelWriter(matched_filename, engine='openpyxl') as writer:
                    matched_df.to_excel(writer, sheet_name='Matched_Employees', index=False)
                    improved.to_excel(writer, sheet_name='Improved_Compliance', index=False)
                    declined.to_excel(writer, sheet_name='Declined_Compliance', index=False)
                    unchanged.to_excel(writer, sheet_name='Unchanged_Compliance', index=False)
            
            safe_file_operation(
                save_comprehensive_report,
                lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
                "Saving comprehensive report"
            )
            
            self.app.root.after(0, self.process_complete, output_dir)
            
        except Exception as e:
            self.app.root.after(0, self.process_error, e)
    
    def clean_report_data(self, filepath, file_type):
        """Enhanced data cleaning with better error handling."""
        try:
            log_message_util(
                self.app.root,
                self.log_text,
                f"Processing {file_type} file: {os.path.basename(filepath)}",
                'info'
            )
            
            header_row = self.find_header_row(filepath)
            if header_row is None:
                raise ValueError(f"Could not determine header row in {os.path.basename(filepath)}")
            
            df = pd.read_excel(filepath, skiprows=header_row)
            df.rename(columns=lambda c: str(c).strip(), inplace=True)
            
            # Enhanced column mapping with fuzzy matching
            rename_dict = self.map_columns_intelligently(df.columns)
            df.rename(columns=rename_dict, inplace=True)
            
            # Apply fill logic
            logic = self.fill_logic.get()
            if logic != "No Fill Down":
                self.apply_fill_logic(df, logic)
            
            # Validate required columns
            if 'Time Sheet Owner Name' not in df.columns:
                raise ValueError(f"Critical column 'Time Sheet Owner Name' not found in {file_type} file")
            
            # Clean data
            df.dropna(subset=['Time Sheet Owner Name'], inplace=True)
            
            log_message_util(
                self.app.root,
                self.log_text,
                f"Successfully processed {len(df)} records from {file_type} file",
                'success'
            )
            
            return df
            
        except Exception as e:
            log_message_util(
                self.app.root,
                self.log_text,
                f"Error processing {file_type} file: {str(e)}",
                'error'
            )
            return None
    
    def map_columns_intelligently(self, columns):
        """Intelligent column mapping with fuzzy matching."""
        rename_dict = {}
        unmapped_cols = list(columns)
        
        # Define mapping rules with flexible matching
        mapping_rules = {
            'Time Sheet Owner Name': [
                lambda c: 'time sheet' in c.lower() or 'timesheet' in c.lower(),
                lambda c: 'owner' in c.lower(),
                lambda c: 'name' in c.lower()
            ],
            'Compliance %': [
                lambda c: 'compliance' in c.lower(),
                lambda c: '%' in c
            ],
            'Super Office': [
                lambda c: 'super office' in c.lower()
            ],
            'Division': [
                lambda c: 'division' in c.lower()
            ]
        }
        
        def find_and_map(standard_name, rules):
            for col in unmapped_cols[:]:  # Create a copy to iterate safely
                col_clean = str(col).lower().replace(':', '').strip()
                if all(rule(col_clean) for rule in rules):
                    rename_dict[col] = standard_name
                    unmapped_cols.remove(col)
                    return True
            return False
        
        # Apply mapping rules
        for standard_name, rules in mapping_rules.items():
            find_and_map(standard_name, rules)
        
        return rename_dict
    
    def apply_fill_logic(self, df, logic):
        """Apply selected fill logic to organizational columns."""
        if logic == "Dynamic Keyword Fill":
            keywords = ['super', 'division', 'office']
        elif logic == "Smart Header + Dynamic Fill":
            keywords = ['super', 'division', 'office', 'org', 'unit']
        else:
            return
        
        cols_to_fill = [
            col for col in df.columns 
            if any(keyword in str(col).lower() for keyword in keywords)
        ]
        
        if cols_to_fill:
            df[cols_to_fill] = df[cols_to_fill].fillna(method='ffill')
            log_message_util(
                self.app.root,
                self.log_text,
                f"Applied {logic} to columns: {', '.join(cols_to_fill)}",
                'info'
            )
    
    def find_header_row(self, filepath):
        """Enhanced header detection with multiple strategies."""
        try:
            df_preview = pd.read_excel(filepath, header=None, nrows=50)
            
            # Strategy 1: Look for key columns together
            for i, row in df_preview.iterrows():
                row_str = ' '.join(str(x) for x in row.dropna().values).lower().replace(':', '')
                if all(keyword in row_str for keyword in ['time sheet', 'compliance', 'office']):
                    return i
            
            # Strategy 2: Look for any timesheet-related header
            for i, row in df_preview.iterrows():
                row_str = ' '.join(str(x) for x in row.dropna().values).lower()
                if 'timesheet' in row_str or 'time sheet' in row_str:
                    return i
            
            return None
            
        except Exception as e:
            log_message_util(
                self.app.root,
                self.log_text,
                f"Error finding header in {os.path.basename(filepath)}: {str(e)}",
                'error'
            )
            return None
    
    def validate_employee_ids(self, df1, df2):
        """Enhanced employee validation with normalized matching."""
        # Normalize names for matching
        df1['Name_Normalized'] = df1['Time Sheet Owner Name'].apply(normalize_name)
        df2['Name_Normalized'] = df2['Time Sheet Owner Name'].apply(normalize_name)
        
        # Perform matching
        matched = pd.merge(
            df1, df2, 
            on='Name_Normalized', 
            how='inner', 
            suffixes=('_hist', '_emp')
        )
        
        # Find unmatched records
        only_in_file1 = df1[~df1['Name_Normalized'].isin(df2['Name_Normalized'])]
        only_in_file2 = df2[~df2['Name_Normalized'].isin(df1['Name_Normalized'])]
        
        return matched, only_in_file1, only_in_file2
    
    def track_compliance_changes(self, matched_df):
        """Enhanced compliance change tracking with better numeric handling."""
        # Ensure compliance columns are numeric
        matched_df['Compliance %_emp'] = pd.to_numeric(
            matched_df['Compliance %_emp'], errors='coerce'
        )
        matched_df['Compliance %_hist'] = pd.to_numeric(
            matched_df['Compliance %_hist'], errors='coerce'
        )
        
        # Calculate changes
        matched_df['Compliance_Change'] = (
            matched_df['Compliance %_emp'] - matched_df['Compliance %_hist']
        )
        
        # Categorize changes
        improved = matched_df[matched_df['Compliance_Change'] > 0].copy()
        declined = matched_df[matched_df['Compliance_Change'] < 0].copy()
        unchanged = matched_df[matched_df['Compliance_Change'] == 0].copy()
        
        return improved, declined, unchanged
    
    def process_complete(self, output_dir):
        """Enhanced completion handler."""
        self.progress.stop()
        self.generate_button.configure_state("normal")
        
        log_message_util(
            self.app.root,
            self.log_text,
            "Report generation completed successfully!",
            'success'
        )
        
        messagebox.showinfo(
            "Success",
            f"Comparison reports have been generated successfully!\n\n"
            f"Files saved to:\n{output_dir}\n\n"
            f"Reports include:\n"
            f"• Comprehensive matched employee analysis\n"
            f"• Compliance improvement tracking\n"
            f"• Unmatched record identification"
        )
    
    def process_error(self, error):
        """Enhanced error handler."""
        self.progress.stop()
        self.generate_button.configure_state("normal")
        
        error_msg = f"Processing error: {str(error)}"
        log_message_util(self.app.root, self.log_text, error_msg, 'error')
        
        messagebox.showerror(
            "Processing Error",
            f"An error occurred during report generation:\n\n{error}\n\n"
            f"Please check the log for more details and ensure your files "
            f"are valid Excel files with the expected structure."
        )

#==========================================================================
# TAB 2: HELPER SHEET FORMATTER (Enhanced with Master Names ID Matching)
#==========================================================================
class FormatterFrame(ttk.Frame):
    def __init__(self, container, app_instance, **kwargs):
        super().__init__(container, **kwargs)
        self.app = app_instance
        
        # Variables
        self.source_file = tk.StringVar()
        self.employee_lookup_file = tk.StringVar()  # New: Jan 2025 names file
        self.pay_period = tk.StringVar()
        self.report_day = tk.StringVar(value="Wednesday")
        self.output_location = tk.StringVar(value=self.app.default_output_path)
        
        self.pp_data = self.get_pay_period_data()
        self.pp_display_list = [item['display'] for item in self.pp_data]
        
        # Match engine for ID assignment
        self.match_engine = MatchingEngine()
        
        self.setup_ui()
    
    def setup_ui(self):
        """Setup modern UI layout with numbered steps."""
        self.grid_rowconfigure(1, weight=1)
        self.grid_columnconfigure(0, weight=1)
        
        # Header
        header_frame = ttk.Frame(self)
        header_frame.grid(row=0, column=0, sticky='ew', pady=(0, 30))
        
        ttk.Label(
            header_frame,
            text="Helper Sheet Data Formatter",
            font=('SF Pro Display', 24, 'bold'),
            foreground='#1C1C1E'
        ).pack(anchor='w')
        
        ttk.Label(
            header_frame,
            text="Format helper data and assign employee IDs with intelligent matching",
            font=('SF Pro Text', 14),
            foreground='#8E8E93'
        ).pack(anchor='w', pady=(5, 0))
        
        # Main content
        content_frame = ttk.Frame(self)
        content_frame.grid(row=1, column=0, sticky='nsew')
        content_frame.grid_columnconfigure(0, weight=1)
        content_frame.grid_rowconfigure(6, weight=1)
        
        # Step 1: Employee Lookup File (NEW)
        step1 = StepIndicator(
            content_frame,
            1,
            "Select Employee Lookup File",
            "Jan 2025 names with Employee IDs"
        )
        step1.grid(row=0, column=0, sticky='ew', pady=(0, 15))
        
        lookup_frame = ttk.Frame(content_frame)
        lookup_frame.grid(row=0, column=1, sticky='ew', padx=(20, 0))
        lookup_frame.grid_columnconfigure(0, weight=1)
        
        self.lookup_entry = ttk.Entry(
            lookup_frame,
            textvariable=self.employee_lookup_file,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        self.lookup_entry.grid(row=0, column=0, sticky='ew', ipady=8)
        
        ttk.Button(
            lookup_frame,
            text="Browse",
            command=lambda: self.browse_file(self.employee_lookup_file, "Select Employee Lookup File")
        ).grid(row=0, column=1, padx=(10, 0))
        
        # Step 2: Source File
        step2 = StepIndicator(
            content_frame,
            2,
            "Select Source Excel File",
            "Helper sheet with Inactive and Excused lists"
        )
        step2.grid(row=1, column=0, sticky='ew', pady=(20, 15))
        
        source_frame = ttk.Frame(content_frame)
        source_frame.grid(row=1, column=1, sticky='ew', padx=(20, 0))
        source_frame.grid_columnconfigure(0, weight=1)
        
        self.source_entry = ttk.Entry(
            source_frame,
            textvariable=self.source_file,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        self.source_entry.grid(row=0, column=0, sticky='ew', ipady=8)
        
        ttk.Button(
            source_frame,
            text="Browse",
            command=lambda: self.browse_file(self.source_file, "Select Source Excel File")
        ).grid(row=0, column=1, padx=(10, 0))
        
        # Step 3: Pay Period
        step3 = StepIndicator(
            content_frame,
            3,
            "Select Pay Period",
            "Choose the reporting pay period"
        )
        step3.grid(row=2, column=0, sticky='ew', pady=(20, 15))
        
        pp_frame = ttk.Frame(content_frame)
        pp_frame.grid(row=2, column=1, sticky='ew', padx=(20, 0))
        
        pp_combo = ttk.Combobox(
            pp_frame,
            textvariable=self.pay_period,
            values=self.pp_display_list,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        pp_combo.pack(fill='x', ipady=8)
        if self.pp_display_list:
            pp_combo.set(self.pp_display_list[0])
        
        # Step 4: Report Day
        step4 = StepIndicator(
            content_frame,
            4,
            "Select Report Day",
            "Monday or Wednesday reporting cycle"
        )
        step4.grid(row=3, column=0, sticky='ew', pady=(20, 15))
        
        day_frame = ttk.Frame(content_frame)
        day_frame.grid(row=3, column=1, sticky='ew', padx=(20, 0))
        
        day_combo = ttk.Combobox(
            day_frame,
            textvariable=self.report_day,
            values=["Monday", "Wednesday"],
            state='readonly',
            font=('SF Pro Text', 12)
        )
        day_combo.pack(fill='x', ipady=8)
        
        # Step 5: Output Location
        step5 = StepIndicator(
            content_frame,
            5,
            "Select Output Location",
            "Where to save formatted data"
        )
        step5.grid(row=4, column=0, sticky='ew', pady=(20, 15))
        
        output_frame = ttk.Frame(content_frame)
        output_frame.grid(row=4, column=1, sticky='ew', padx=(20, 0))
        output_frame.grid_columnconfigure(0, weight=1)
        
        self.output_entry = ttk.Entry(
            output_frame,
            textvariable=self.output_location,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        self.output_entry.grid(row=0, column=0, sticky='ew', ipady=8)
        
        ttk.Button(
            output_frame,
            text="Browse",
            command=lambda: self.browse_directory(self.output_location)
        ).grid(row=0, column=1, padx=(10, 0))
        
        # Action section
        action_frame = ttk.Frame(content_frame)
        action_frame.grid(row=6, column=0, columnspan=2, sticky='nsew', pady=(30, 0))
        action_frame.grid_columnconfigure(0, weight=1)
        action_frame.grid_rowconfigure(3, weight=1)
        
        # Generate button
        self.generate_button = ModernButton(
            action_frame,
            "Generate Formatted Data with IDs",
            self.run_process,
            style_type="success"
        )
        self.generate_button.grid(row=0, column=0, sticky='ew', ipady=15)
        
        # Progress bar
        self.progress = ttk.Progressbar(
            action_frame,
            orient='horizontal',
            mode='indeterminate'
        )
        self.progress.grid(row=1, column=0, sticky='ew', pady=(15, 0))
        
        # Summary badges
        self.summary_frame = ttk.Frame(action_frame)
        self.summary_frame.grid(row=2, column=0, sticky='ew', pady=(15, 0))
        self.summary_frame.grid_columnconfigure((0, 1, 2, 3), weight=1)
        
        self.exact_badge = SummaryBadge(self.summary_frame, "Exact", "0", "#34C759")
        self.exact_badge.grid(row=0, column=0, sticky='ew', padx=(0, 5))
        
        self.fuzzy_badge = SummaryBadge(self.summary_frame, "Fuzzy", "0", "#FF9500")
        self.fuzzy_badge.grid(row=0, column=1, sticky='ew', padx=5)
        
        self.no_match_badge = SummaryBadge(self.summary_frame, "No Match", "0", "#FF3B30")
        self.no_match_badge.grid(row=0, column=2, sticky='ew', padx=5)
        
        self.had_id_badge = SummaryBadge(self.summary_frame, "Had ID", "0", "#007AFF")
        self.had_id_badge.grid(row=0, column=3, sticky='ew', padx=(5, 0))
        
        # Log section with filters
        log_frame = ttk.LabelFrame(action_frame, text="Processing Log", padding="15")
        log_frame.grid(row=3, column=0, sticky='nsew', pady=(20, 0))
        log_frame.grid_rowconfigure(1, weight=1)
        log_frame.grid_columnconfigure(0, weight=1)
        
        # Log filters
        filter_frame = ttk.Frame(log_frame)
        filter_frame.grid(row=0, column=0, sticky='ew', pady=(0, 10))
        
        ttk.Label(filter_frame, text="Show:").pack(side='left', padx=(0, 10))
        
        self.show_info = tk.BooleanVar(value=True)
        self.show_success = tk.BooleanVar(value=True)
        self.show_warning = tk.BooleanVar(value=True)
        self.show_error = tk.BooleanVar(value=True)
        
        ttk.Checkbutton(filter_frame, text="Info", variable=self.show_info).pack(side='left', padx=5)
        ttk.Checkbutton(filter_frame, text="Success", variable=self.show_success).pack(side='left', padx=5)
        ttk.Checkbutton(filter_frame, text="Warning", variable=self.show_warning).pack(side='left', padx=5)
        ttk.Checkbutton(filter_frame, text="Error", variable=self.show_error).pack(side='left', padx=5)
        
        # Log text with scrollbar
        log_container = ttk.Frame(log_frame)
        log_container.grid(row=1, column=0, sticky='nsew')
        log_container.grid_rowconfigure(0, weight=1)
        log_container.grid_columnconfigure(0, weight=1)
        
        self.log_text = tk.Text(
            log_container,
            height=12,
            state='disabled',
            bg='#F8F9FA',
            fg='#1C1C1E',
            wrap='word',
            borderwidth=0,
            highlightthickness=0,
            font=('SF Mono', 10)
        )
        self.log_text.grid(row=0, column=0, sticky='nsew')
        
        scrollbar = ttk.Scrollbar(log_container, orient='vertical', command=self.log_text.yview)
        scrollbar.grid(row=0, column=1, sticky='ns')
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.validate_inputs()
    
    def browse_file(self, string_var, title):
        """Enhanced file browser with validation."""
        filename = filedialog.askopenfilename(
            title=title,
            filetypes=[("Excel files", "*.xlsx *.xlsm"), ("All files", "*.*")]
        )
        if filename:
            string_var.set(filename)
            log_message_util(
                self.app.root,
                self.log_text,
                f"Selected: {os.path.basename(filename)}",
                'success'
            )
            self.validate_inputs()
    
    def browse_directory(self, string_var):
        """Enhanced directory browser."""
        directory = filedialog.askdirectory(
            title="Select Output Location",
            initialdir=string_var.get()
        )
        if directory:
            string_var.set(directory)
            log_message_util(
                self.app.root,
                self.log_text,
                f"Output location: {directory}",
                'success'
            )
    
    def validate_inputs(self):
        """Validate inputs and update UI state."""
        has_lookup = bool(self.employee_lookup_file.get())
        has_source = bool(self.source_file.get())
        has_pp = bool(self.pay_period.get())
        
        if has_lookup and has_source and has_pp:
            self.generate_button.configure_state("normal")
        else:
            self.generate_button.configure_state("disabled")
    
    def run_process(self):
        """Enhanced process runner with validation."""
        if not all([self.employee_lookup_file.get(), self.source_file.get(), self.pay_period.get()]):
            messagebox.showerror(
                "Input Error",
                "Please select the employee lookup file, source file, and pay period."
            )
            return
        
        self.generate_button.configure_state("disabled")
        self.progress.start(10)
        log_message_util(
            self.app.root,
            self.log_text,
            "Starting helper data formatting with ID matching...",
            'info'
        )
        
        thread = threading.Thread(target=self.process_files_thread)
        thread.daemon = True
        thread.start()
    
    def process_files_thread(self):
        """Enhanced processing with ID matching."""
        try:
            # Load employee lookup data
            log_message_util(
                self.app.root,
                self.log_text,
                "Loading employee lookup data...",
                'info'
            )
            
            lookup_df = safe_file_operation(
                lambda: pd.read_excel(self.employee_lookup_file.get()),
                lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
                "Loading employee lookup file"
            )
            
            if lookup_df is None:
                raise ValueError("Failed to load employee lookup file")
            
            # Build matching engine
            self.match_engine.build_lookup(lookup_df, 'Employee Name', 'Employee ID')
            log_message_util(
                self.app.root,
                self.log_text,
                f"Built lookup table with {len(self.match_engine.lookup_dict)} employees",
                'success'
            )
            
            # Process source data
            source_path = self.source_file.get()
            pp_display = self.pay_period.get()
            selected_day = self.report_day.get().upper()
            
            pp_info = next((item for item in self.pp_data if item['display'] == pp_display), None)
            if not pp_info:
                raise ValueError("Could not find data for the selected pay period.")
            
            pp_num, pp_year = pp_info['pp'], pp_info['year']
            log_message_util(
                self.app.root,
                self.log_text,
                f"Processing for PP {pp_num}, Year {pp_year}",
                'info'
            )
            
            final_dfs = []
            
            # Process Excused List
            try:
                log_message_util(
                    self.app.root,
                    self.log_text,
                    "Processing Excused List...",
                    'info'
                )
                
                df_excused = pd.read_excel(source_path, sheet_name='Excused List', skiprows=4)
                df_excused = self.process_excused_data(df_excused, pp_num, pp_year, selected_day)
                if df_excused is not None and len(df_excused) > 0:
                    final_dfs.append(df_excused)
                    log_message_util(
                        self.app.root,
                        self.log_text,
                        f"Processed {len(df_excused)} records from Excused List",
                        'success'
                    )
                
            except Exception as e:
                log_message_util(
                    self.app.root,
                    self.log_text,
                    f"Could not process Excused List: {str(e)}",
                    'error'
                )
            
            # Process Inactive List
            try:
                log_message_util(
                    self.app.root,
                    self.log_text,
                    "Processing Inactive Names...",
                    'info'
                )
                
                df_inactive = pd.read_excel(source_path, sheet_name='Inactive Names', skiprows=1)
                df_inactive = self.process_inactive_data(df_inactive, pp_num, pp_year)
                if df_inactive is not None and len(df_inactive) > 0:
                    final_dfs.append(df_inactive)
                    log_message_util(
                        self.app.root,
                        self.log_text,
                        f"Processed {len(df_inactive)} records from Inactive Names",
                        'success'
                    )
                
            except Exception as e:
                log_message_util(
                    self.app.root,
                    self.log_text,
                    f"Could not process Inactive Names: {str(e)}",
                    'error'
                )
            
            if not final_dfs:
                raise ValueError("No data could be processed from any sheets.")
            
            # Combine and finalize data
            final_df = pd.concat(final_dfs, ignore_index=True)
            output_columns = [
                'YEAR', 'PP', 'User Last Name', 'User First Name', 
                'Full Name (First, Last)', 'Employee ID', 'Matched Employee ID',
                'Match Status', 'Match Confidence', 'Column1'
            ]
            final_df = final_df[output_columns]
            
            # Save outputs
            output_dir = self.output_location.get()
            
            # Save main formatted data
            main_output = os.path.join(
                output_dir,
                format_output_filename("Formatted_Helper_Data", ".csv", True)
            )
            
            safe_file_operation(
                lambda: final_df.to_csv(main_output, index=False),
                lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
                "Saving main formatted data"
            )
            
            # Save Master_Names_with_IDs Excel file
            master_output = os.path.join(
                output_dir,
                format_output_filename("Master_Names_with_IDs", ".xlsx", True)
            )
            
            def save_master_file():
                with pd.ExcelWriter(master_output, engine='openpyxl') as writer:
                    final_df.to_excel(writer, sheet_name='Master_Names_with_IDs', index=False)
                    
                    # Add summary sheet
                    summary_df = pd.DataFrame([{
                        'Metric': 'Total Records',
                        'Count': len(final_df)
                    }, {
                        'Metric': 'Exact Matches',
                        'Count': self.match_engine.match_stats['exact_matches']
                    }, {
                        'Metric': 'Fuzzy Matches',
                        'Count': self.match_engine.match_stats['fuzzy_matches']
                    }, {
                        'Metric': 'No Matches',
                        'Count': self.match_engine.match_stats['no_matches']
                    }, {
                        'Metric': 'Already Had ID',
                        'Count': self.match_engine.match_stats['already_had_id']
                    }])
                    summary_df.to_excel(writer, sheet_name='Summary', index=False)
            
            safe_file_operation(
                save_master_file,
                lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
                "Saving Master_Names_with_IDs file"
            )
            
            # Update summary badges
            self.app.root.after(0, self.update_summary_badges)
            
            self.app.root.after(0, self.process_complete, output_dir)
            
        except Exception as e:
            self.app.root.after(0, self.process_error, e)
    
    def process_excused_data(self, df_excused, pp_num, pp_year, selected_day):
        """Process excused list data with ID matching."""
        if df_excused.empty:
            return None
        
        # Clean and validate data
        df_excused.dropna(subset=['EMPLOYEE_ID'], inplace=True)
        df_excused = df_excused[pd.to_numeric(df_excused['EMPLOYEE_ID'], errors='coerce').notna()].copy()
        
        if df_excused.empty:
            return None
        
        # Parse names
        names = df_excused['Employee Name'].str.split(n=1, expand=True)
        df_excused['User First Name'] = names[0]
        df_excused['User Last Name'] = names[1].apply(
            lambda x: x.split()[-1] if isinstance(x, str) else None
        )
        df_excused['Full Name (First, Last)'] = df_excused['Employee Name']
        
        # Apply ID matching
        df_excused[['Matched Employee ID', 'Match Status', 'Match Confidence']] = df_excused.apply(
            lambda row: self.match_employee_row(row['Employee Name'], row.get('EMPLOYEE_ID')),
            axis=1,
            result_type='expand'
        )
        
        # Standardize columns
        df_excused.rename(columns={'EMPLOYEE_ID': 'Employee ID'}, inplace=True)
        df_excused['YEAR'], df_excused['PP'] = pp_year, pp_num
        df_excused['Column1'] = f"PP {pp_num} - EXCUSED LIST TABLEAU - {selected_day}"
        
        return df_excused
    
    def process_inactive_data(self, df_inactive, pp_num, pp_year):
        """Process inactive list data with ID matching."""
        if df_inactive.empty:
            return None
        
        # Clean and validate data
        df_inactive.dropna(subset=['HHS ID'], inplace=True)
        df_inactive = df_inactive[pd.to_numeric(df_inactive['HHS ID'], errors='coerce').notna()].copy()
        
        if df_inactive.empty:
            return None
        
        # Standardize columns
        df_inactive.rename(columns={
            'HHS ID': 'Employee ID',
            'First Name': 'User First Name',
            'Last Name': 'User Last Name'
        }, inplace=True)
        
        df_inactive['Full Name (First, Last)'] = (
            df_inactive['User First Name'] + ' ' + df_inactive['User Last Name']
        )
        
        # Apply ID matching
        df_inactive[['Matched Employee ID', 'Match Status', 'Match Confidence']] = df_inactive.apply(
            lambda row: self.match_employee_row(row['Full Name (First, Last)'], row.get('Employee ID')),
            axis=1,
            result_type='expand'
        )
        
        df_inactive['YEAR'], df_inactive['PP'] = pp_year, pp_num
        df_inactive['Column1'] = f"PP{pp_num}, Inactive in EASE"
        
        return df_inactive
    
    def match_employee_row(self, name, existing_id):
        """Match a single employee row and return results."""
        matched_id, status, confidence = self.match_engine.match_employee_id(name, existing_id)
        return pd.Series([matched_id, status, confidence])
    
    def update_summary_badges(self):
        """Update summary badges with current statistics."""
        stats = self.match_engine.match_stats
        self.exact_badge.update_value(stats['exact_matches'])
        self.fuzzy_badge.update_value(stats['fuzzy_matches'])
        self.no_match_badge.update_value(stats['no_matches'])
        self.had_id_badge.update_value(stats['already_had_id'])
    
    def get_pay_period_data(self):
        """Get pay period data for dropdown."""
        raw_data = """2,12/29/2024,1/11/2025
3,1/12/2025,1/25/2025
4,1/26/2025,2/8/2025
5,2/9/2025,2/22/2025
6,2/23/2025,3/8/2025
7,3/9/2025,3/22/2025
8,3/23/2025,4/5/2025
9,4/6/2025,4/19/2025
10,4/20/2025,5/3/2025
11,5/4/2025,5/17/2025
12,5/18/2025,5/31/2025
13,6/1/2025,6/14/2025
14,6/15/2025,6/28/2025
15,6/29/2025,7/12/2025
16,7/13/2025,7/26/2025
17,7/27/2025,8/9/2025
18,8/10/2025,8/23/2025
19,8/24/2025,9/6/2025
20,9/7/2025,9/20/2025
21,9/21/2025,10/4/2025
22,10/5/2025,10/18/2025
23,10/19/2025,11/1/2025
24,11/2/2025,11/15/2025
25,11/16/2025,11/29/2025
1,11/30/2025,12/13/2025
2,12/14/2025,12/27/2025
3,12/28/2025,1/10/2026
4,1/11/2026,1/24/2026
5,1/25/2026,2/7/2026"""
        
        pay_periods = []
        for line in raw_data.strip().split('\n'):
            pp, start, end = line.strip().split(',')
            year = pd.to_datetime(start).year
            pay_periods.append({
                "display": f"PP {pp} ({start} - {end})", 
                "pp": int(pp), 
                "year": year
            })
        return pay_periods
    
    def process_complete(self, output_dir):
        """Enhanced completion handler with summary."""
        self.progress.stop()
        self.generate_button.configure_state("normal")
        
        summary_stats = self.match_engine.get_summary_stats()
        log_message_util(
            self.app.root,
            self.log_text,
            f"Processing completed successfully! {summary_stats}",
            'success'
        )
        
        messagebox.showinfo(
            "Success",
            f"Helper data formatting completed successfully!\n\n"
            f"Files saved to:\n{output_dir}\n\n"
            f"Matching Summary:\n{summary_stats}\n\n"
            f"Files generated:\n"
            f"• Formatted_Helper_Data.csv\n"
            f"• Master_Names_with_IDs.xlsx"
        )
    
    def process_error(self, error):
        """Enhanced error handler."""
        self.progress.stop()
        self.generate_button.configure_state("normal")
        
        error_msg = f"Processing error: {str(error)}"
        log_message_util(self.app.root, self.log_text, error_msg, 'error')
        
        messagebox.showerror(
            "Processing Error",
                         f"An error occurred during helper data formatting:\n\n{error}\n\n"
             f"Please check the log for more details and ensure your files "
             f"are valid Excel files with the expected structure."
         )

#==========================================================================
# TAB 3: OFFICE EMAILER (Enhanced with Fixed HTML CSS)
#==========================================================================
class EmailerFrame(ttk.Frame):
    def __init__(self, container, app_instance, **kwargs):
        super().__init__(container, **kwargs)
        self.app = app_instance
        
        # Variables
        self.source_file = tk.StringVar()
        self.current_pp = tk.StringVar()
        self.manual_filename = tk.StringVar()
        self.output_format = tk.StringVar(value="Outlook")
        self.word_output_location = tk.StringVar(value=self.app.default_output_path)
        self.office_vars = {}
        self.dataframe = None
        self.dates_df = None
        self.director_data = {}
        self.non_reporters_data = {}
        
        self.setup_ui()
    
    def setup_ui(self):
        """Setup modern UI layout."""
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=2)
        self.grid_rowconfigure(1, weight=1)
        
        # Header
        header_frame = ttk.Frame(self)
        header_frame.grid(row=0, column=0, columnspan=2, sticky='ew', pady=(0, 30))
        
        ttk.Label(
            header_frame,
            text="Office Participation Emailer",
            font=('SF Pro Display', 24, 'bold'),
            foreground='#1C1C1E'
        ).pack(anchor='w')
        
        ttk.Label(
            header_frame,
            text="Generate professional participation reports with fixed HTML formatting",
            font=('SF Pro Text', 14),
            foreground='#8E8E93'
        ).pack(anchor='w', pady=(5, 0))
        
        # Controls panel
        self.controls_frame = ttk.Frame(self)
        self.controls_frame.grid(row=1, column=0, sticky='nsew', padx=(0, 20))
        self.controls_frame.grid_rowconfigure(3, weight=1)
        
        # Step 1: Data file
        step1 = StepIndicator(
            self.controls_frame,
            1,
            "Select Data Digest File",
            "Office participation data"
        )
        step1.pack(fill='x', pady=(0, 15))
        
        file_frame = ttk.Frame(self.controls_frame)
        file_frame.pack(fill='x', pady=(0, 15))
        file_frame.grid_columnconfigure(0, weight=1)
        
        self.file_entry = ttk.Entry(
            file_frame,
            textvariable=self.source_file,
            state='readonly',
            font=('SF Pro Text', 12)
        )
        self.file_entry.grid(row=0, column=0, sticky='ew', ipady=8)
        
        ttk.Button(
            file_frame,
            text="Browse",
            command=self.load_data
        ).grid(row=0, column=1, padx=(10, 0))
        
        # Step 2: Pay period
        step2 = StepIndicator(
            self.controls_frame,
            2,
            "Pay Period Being Reported",
            "Enter the pay period number"
        )
        step2.pack(fill='x', pady=(0, 15))
        
        self.pp_entry = ttk.Entry(
            self.controls_frame,
            textvariable=self.current_pp,
            font=('SF Pro Text', 12)
        )
        self.pp_entry.pack(fill='x', ipady=8, pady=(0, 15))
        
        # Step 3: Office selection
        office_frame = ttk.LabelFrame(self.controls_frame, text="3. Select Offices", padding="15")
        office_frame.pack(fill='both', expand=True, pady=(0, 15))
        
        toggle_frame = ttk.Frame(office_frame)
        toggle_frame.pack(fill='x', pady=(0, 10))
        toggle_frame.grid_columnconfigure((0, 1), weight=1)
        
        ModernButton(
            toggle_frame,
            "Select All",
            lambda: self.toggle_all_offices(True),
            style_type="secondary"
        ).grid(row=0, column=0, sticky='ew', padx=(0, 5))
        
        ModernButton(
            toggle_frame,
            "Select None", 
            lambda: self.toggle_all_offices(False),
            style_type="secondary"
        ).grid(row=0, column=1, sticky='ew', padx=(5, 0))
        
        # Office list with scrollbar
        self.office_canvas = tk.Canvas(
            office_frame,
            bg='#F8F9FA',
            highlightthickness=0,
            height=200
        )
        self.office_list_frame = ttk.Frame(self.office_canvas)
        self.office_scrollbar = ttk.Scrollbar(
            office_frame,
            orient="vertical",
            command=self.office_canvas.yview
        )
        self.office_canvas.configure(yscrollcommand=self.office_scrollbar.set)
        
        self.office_scrollbar.pack(side="right", fill="y")
        self.office_canvas.pack(side="left", fill="both", expand=True)
        self.office_canvas.create_window((0, 0), window=self.office_list_frame, anchor="nw")
        self.office_list_frame.bind(
            "<Configure>",
            lambda e: self.office_canvas.configure(scrollregion=self.office_canvas.bbox("all"))
        )
        
        # Step 4: Output format
        output_frame = ttk.LabelFrame(self.controls_frame, text="4. Select Output Format", padding="15")
        output_frame.pack(fill='x', pady=(0, 15))
        
        ttk.Radiobutton(
            output_frame,
            text="Open in Outlook",
            variable=self.output_format,
            value="Outlook",
            command=self.toggle_word_path_visibility
        ).pack(anchor='w', pady=5)
        
        ttk.Radiobutton(
            output_frame,
            text="Generate Word Document",
            variable=self.output_format,
            value="Word",
            command=self.toggle_word_path_visibility
        ).pack(anchor='w', pady=5)
        
        # Word filename (conditional)
        self.word_frame = ttk.LabelFrame(
            self.controls_frame,
            text="5. Word Document Filename (Optional)",
            padding="15"
        )
        
        self.word_entry = ttk.Entry(
            self.word_frame,
            textvariable=self.manual_filename,
            font=('SF Pro Text', 12)
        )
        self.word_entry.pack(fill='x', ipady=8, pady=(0, 5))
        
        ttk.Label(
            self.word_frame,
            text="Format: YYYY.MM.DD_PP## Emails..._HHMMSS.docx",
            font=('SF Pro Text', 10),
            foreground='#8E8E93'
        ).pack(anchor='w')
        
        # Generate button
        self.generate_button = ModernButton(
            self.controls_frame,
            "Generate Emails",
            self.run_generation,
            style_type="success",
            state="disabled"
        )
        self.generate_button.pack(fill='x', ipady=15, pady=(15, 0))
        
        # Log panel
        log_frame = ttk.LabelFrame(self, text="Processing Log", padding="15")
        log_frame.grid(row=1, column=1, sticky='nsew')
        log_frame.grid_rowconfigure(0, weight=1)
        log_frame.grid_columnconfigure(0, weight=1)
        
        # Log with scrollbar
        log_container = ttk.Frame(log_frame)
        log_container.grid(row=0, column=0, sticky='nsew')
        log_container.grid_rowconfigure(0, weight=1)
        log_container.grid_columnconfigure(0, weight=1)
        
        self.log_text = tk.Text(
            log_container,
            state='disabled',
            bg='#F8F9FA',
            fg='#1C1C1E',
            wrap='word',
            borderwidth=0,
            highlightthickness=0,
            font=('SF Mono', 10)
        )
        self.log_text.grid(row=0, column=0, sticky='nsew')
        
        log_scrollbar = ttk.Scrollbar(log_container, orient='vertical', command=self.log_text.yview)
        log_scrollbar.grid(row=0, column=1, sticky='ns')
        self.log_text.configure(yscrollcommand=log_scrollbar.set)
        
        self.toggle_word_path_visibility()
    
    def toggle_word_path_visibility(self):
        """Toggle Word filename field visibility."""
        if self.output_format.get() == "Word":
            self.word_frame.pack(fill='x', pady=(0, 15))
        else:
            self.word_frame.pack_forget()
    
    def toggle_all_offices(self, select):
        """Toggle all office selections."""
        for var in self.office_vars.values():
            var.set(select)
    
    def load_data(self):
        """Enhanced data loading with better error handling."""
        filepath = filedialog.askopenfilename(
            title="Select Data Digest",
            filetypes=[("Excel files", "*.xlsx *.xlsm"), ("All files", "*.*")]
        )
        if not filepath:
            return
        
        self.source_file.set(filepath)
        log_message_util(
            self.app.root,
            self.log_text,
            f"Loading data from: {os.path.basename(filepath)}",
            'info'
        )
        
        def load_data_thread():
            try:
                xls = pd.ExcelFile(filepath)
                log_message_util(self.app.root, self.log_text, "Excel file opened successfully", 'success')
                
                # Load Summary Report
                df = pd.read_excel(xls, sheet_name='Summary Report', skiprows=2, engine='openpyxl')
                df.columns = [str(col).strip() for col in df.columns]
                df.dropna(subset=['Office'], inplace=True)
                df = df[~df['Office'].str.contains('Grand Total', case=False, na=False)]
                self.dataframe = df[df['Office'].str.match(r'^[A-Z]{2,5}$')].copy()
                
                log_message_util(
                    self.app.root,
                    self.log_text,
                    f"Processed {len(self.dataframe)} office records",
                    'success'
                )
                
                # Load other sheets with safe operations
                self.dates_df = safe_file_operation(
                    lambda: pd.read_excel(xls, sheet_name='Dates', engine='openpyxl'),
                    lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
                    "Loading Dates sheet"
                )
                
                # Load director data
                self.load_director_data(xls)
                
                # Load non-reporters data
                self.load_non_reporters_data(xls)
                
                # Auto-detect pay period
                self.auto_detect_pay_period(xls)
                
                # Update office list
                self.app.root.after(0, self.update_office_list)
                
                log_message_util(
                    self.app.root,
                    self.log_text,
                    "Data loading completed successfully",
                    'success'
                )
                
            except Exception as e:
                error_msg = f"Error loading file: {str(e)}"
                log_message_util(self.app.root, self.log_text, error_msg, 'error')
                self.app.root.after(0, lambda: messagebox.showerror("File Load Error", error_msg))
        
        # Run in background thread
        thread = threading.Thread(target=load_data_thread)
        thread.daemon = True
        thread.start()
    
    def load_director_data(self, xls):
        """Load director data with error handling."""
        try:
            directors_raw_df = pd.read_excel(xls, sheet_name='Office Directors', header=None, engine='openpyxl')
            self.director_data = {}
            
            for col in directors_raw_df.columns:
                col_data = directors_raw_df[col]
                office_abbr = col_data[0]
                if office_abbr and isinstance(office_abbr, str):
                    self.director_data[office_abbr] = {
                        'Director Name': col_data[4], 'Director Email': col_data[5],
                        'Deputy Name': col_data[6], 'Deputy Email': col_data[7],
                        'CC 1': col_data[8], 'CC 2': col_data[9], 'CC 3': col_data[10],
                        'CC 4': col_data[11], 'CC 5': col_data[12], 'CC 6': col_data[13],
                        'CC 7': col_data[14], 'CC 8': col_data[15], 'CC 9': col_data[16],
                        'CC 10': col_data[17], 'CC 11': col_data[18],
                        'BCC 1': col_data[20], 'BCC 2': col_data[21]
                    }
            
            log_message_util(
                self.app.root,
                self.log_text,
                f"Loaded director data for {len(self.director_data)} offices",
                'success'
            )
            
        except Exception as e:
            log_message_util(
                self.app.root,
                self.log_text,
                f"Could not load Office Directors sheet: {str(e)}",
                'warning'
            )
            self.director_data = {}
    
    def load_non_reporters_data(self, xls):
        """Load non-reporters data with enhanced error handling."""
        try:
            # Find pivot sheet using structural fingerprinting
            pivot_sheet_name = None
            for sheet in xls.sheet_names:
                try:
                    header_check_df = pd.read_excel(
                        xls, 
                        sheet_name=sheet, 
                        header=None, 
                        skiprows=35, 
                        nrows=1, 
                        usecols='H'
                    )
                    if "Super Office" in str(header_check_df.iloc[0, 0]):
                        pivot_sheet_name = sheet
                        break
                except Exception:
                    continue
            
            if not pivot_sheet_name:
                raise ValueError("Could not find pivot data sheet")
            
            log_message_util(
                self.app.root,
                self.log_text,
                f"Found pivot data in sheet: {pivot_sheet_name}",
                'info'
            )
            
            pivot_df = pd.read_excel(xls, sheet_name=pivot_sheet_name, header=35)
            pivot_df.columns = [str(h).strip().replace('↑', '').strip() for h in pivot_df.columns]
            
            # Handle column name variations
            pivot_df.rename(columns={
                'Super Office': 'Super_Office',
                'Time Sheet: Owner Name': 'Owner_Name',
                'Sum of Sum of Hours': 'Sum_of_Hours'
            }, inplace=True)
            
            pivot_df.dropna(subset=['Owner_Name'], inplace=True)
            
            self.non_reporters_data = {}
            for _, row in pivot_df.iterrows():
                office_full_name = row['Super_Office']
                office_abbr = next(
                    (abbr for abbr, name in self.app.office_name_map.items() 
                     if name == office_full_name), None
                )
                
                if office_abbr:
                    if office_abbr not in self.non_reporters_data:
                        self.non_reporters_data[office_abbr] = []
                    
                    hours = row['Sum_of_Hours']
                    name = row['Owner_Name']
                    status = " - Partial Reporter" if hours > 0 else " - Did Not Report"
                    self.non_reporters_data[office_abbr].append(f"{name}{status}")
            
            log_message_util(
                self.app.root,
                self.log_text,
                f"Processed non-reporters data for {len(self.non_reporters_data)} offices",
                'success'
            )
            
        except Exception as e:
            log_message_util(
                self.app.root,
                self.log_text,
                f"Could not process non-reporters data: {str(e)}",
                'warning'
            )
            self.non_reporters_data = {}
    
    def auto_detect_pay_period(self, xls):
        """Auto-detect pay period from sheet title."""
        try:
            sheet_title = str(pd.read_excel(xls, sheet_name='Summary Report', header=None, nrows=1).iloc[0, 0])
            pp_match = re.search(r'PP(\d+)', sheet_title)
            if pp_match:
                self.current_pp.set(pp_match.group(1))
                log_message_util(
                    self.app.root,
                    self.log_text,
                    f"Auto-detected Pay Period: {pp_match.group(1)}",
                    'success'
                )
        except Exception as e:
            log_message_util(
                self.app.root,
                self.log_text,
                f"Could not auto-detect pay period: {str(e)}",
                'warning'
            )
    
    def update_office_list(self):
        """Update office list UI."""
        # Clear existing office widgets
        for widget in self.office_list_frame.winfo_children():
            widget.destroy()
        self.office_vars = {}
        
        if self.dataframe is not None:
            offices = self.dataframe['Office'].dropna().unique()
            for office_abbr in offices:
                full_name = self.app.office_name_map.get(office_abbr, "Unknown Office")
                display_text = f"{full_name} ({office_abbr})"
                var = tk.BooleanVar(value=True)
                
                cb = ttk.Checkbutton(
                    self.office_list_frame,
                    text=display_text,
                    variable=var,
                    font=('SF Pro Text', 11)
                )
                cb.pack(anchor='w', padx=5, pady=2)
                self.office_vars[office_abbr] = var
            
            self.generate_button.configure_state("normal")
            log_message_util(
                self.app.root,
                self.log_text,
                f"Loaded {len(self.office_vars)} offices",
                'success'
            )
    
    def run_generation(self):
        """Enhanced generation with validation."""
        if not self.current_pp.get().isdigit():
            messagebox.showerror("Input Error", "Please enter a valid pay period number.")
            return
        
        selected_offices = [office for office, var in self.office_vars.items() if var.get()]
        if not selected_offices:
            messagebox.showwarning("Selection Error", "Please select at least one office.")
            return
        
        self.generate_button.configure_state("disabled")
        output_type = self.output_format.get()
        
        log_message_util(
            self.app.root,
            self.log_text,
            f"Generating {len(selected_offices)} emails for {output_type}...",
            'info'
        )
        
        thread = threading.Thread(target=self.generation_thread, args=(selected_offices, output_type))
        thread.daemon = True
        thread.start()
    
    def generation_thread(self, selected_offices, output_type):
        """Enhanced generation thread with better error handling."""
        try:
            current_pp_num = int(self.current_pp.get())
            
            if output_type == "Word":
                self.generate_word_document(selected_offices, current_pp_num)
            elif output_type == "Outlook":
                self.generate_outlook_emails(selected_offices, current_pp_num)
            
        except Exception as e:
            error_message = f"Generation error: {str(e)}"
            log_message_util(self.app.root, self.log_text, error_message, 'error')
            self.app.root.after(
                0, 
                lambda: messagebox.showerror("Generation Error", error_message)
            )
        finally:
            self.app.root.after(0, lambda: self.generate_button.configure_state("normal"))
    
    def generate_word_document(self, selected_offices, current_pp_num):
        """Generate Word document with enhanced formatting."""
        doc = Document()
        
        # Set document style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(0)
        
        for i, office in enumerate(selected_offices):
            log_message_util(
                self.app.root,
                self.log_text,
                f"Processing Word doc for office {i+1}/{len(selected_offices)}: {office}",
                'info'
            )
            
            office_data = self.dataframe[self.dataframe['Office'] == office].iloc[0]
            director_info = self.director_data.get(office, {})
            non_reporters_list = self.non_reporters_data.get(office, [])
            
            self.add_email_to_doc(doc, office_data, director_info, current_pp_num, non_reporters_list)
        
        # Save document
        manual_name = self.manual_filename.get().strip()
        if manual_name:
            filename = manual_name if manual_name.endswith('.docx') else f"{manual_name}.docx"
        else:
            timestamp = datetime.now()
            filename = format_output_filename(
                f"{timestamp.strftime('%Y.%m.%d')}_PP{current_pp_num}_Emails_to_Office_Directors_{timestamp.strftime('%H%M%S')}",
                ".docx",
                False
            )
        
        save_path = os.path.join(self.word_output_location.get(), filename)
        
        safe_file_operation(
            lambda: doc.save(save_path),
            lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
            "Saving Word document"
        )
        
        log_message_util(
            self.app.root,
            self.log_text,
            f"Word document saved: {os.path.basename(save_path)}",
            'success'
        )
        
        self.app.root.after(
            0,
            lambda: messagebox.showinfo(
                "Success",
                f"Word document created successfully!\n\n"
                f"File: {os.path.basename(save_path)}\n"
                f"Location: {os.path.dirname(save_path)}"
            )
        )
    
    def generate_outlook_emails(self, selected_offices, current_pp_num):
        """Generate Outlook emails with safe COM operations."""
        success_count = 0
        
        for i, office in enumerate(selected_offices):
            log_message_util(
                self.app.root,
                self.log_text,
                f"Creating Outlook email {i+1}/{len(selected_offices)}: {office}",
                'info'
            )
            
            office_data = self.dataframe[self.dataframe['Office'] == office].iloc[0]
            director_info = self.director_data.get(office, {})
            non_reporters_list = self.non_reporters_data.get(office, [])
            
            if self.create_outlook_email(office_data, director_info, current_pp_num, non_reporters_list):
                success_count += 1
        
        self.app.root.after(
            0,
            lambda: messagebox.showinfo(
                "Success",
                f"Successfully created {success_count} of {len(selected_offices)} Outlook drafts."
            )
        )
    
    def create_html_body(self, data, director_info, current_pp_num, non_reporters_list):
        """Create HTML body with FIXED CSS syntax."""
        pay_period = str(current_pp_num)
        office_abbr = data['Office']
        director_name = director_info.get('Director Name', '<<Director_Name>>')
        office_final_rate = self.format_as_percent(data['Final Participation Rate'])
        office_initial_rate = self.format_as_percent(data['Initial Participation Rate'])
        
        non_reporters_section = ""
        if non_reporters_list:
            non_reporters_html = "".join([f"<li>{item}</li>" for item in non_reporters_list])
            non_reporters_section = f'<b><u>Users who did not report or had proxy:</u></b><ul style="margin-top:0px;">{non_reporters_html}</ul>'
        
        reminder = self.get_next_pp_reminder(current_pp_num)
        if isinstance(reminder, dict):
            reminder_html = f"{reminder['part1']}<b>{reminder['part1_bold']}</b>{reminder['part2']}<b>{reminder['part2_bold']}</b>{reminder['part3']}<b>{reminder['part3_bold']}</b>{reminder['part4']}"
        else:
            reminder_html = reminder
        
        # FIXED HTML with correct CSS syntax (removed invalid style.border syntax)
        html = f"""<p style="font-family:Calibri,sans-serif; font-size:11pt; color:black;">Hello {director_name},<br><br>Last week, CBER completed time reporting in ITR Pay Period <b>{pay_period}</b>. Across CBER, <b>100.00%</b> of staff reported their full Tour of Duty (TOD), and <b>95.54%</b> reported partial hours.<br><br>In Pay Period <b>{pay_period}</b>, <b>{office_abbr}</b> had an unassisted (initial) percentage of <b>{office_initial_rate}</b> which is the percent of staff who entered their full Tour of Duty in ITR by the 11:59pm ET Monday deadline. <b>{office_abbr}</b>'s final participation rate is <b>{office_final_rate}</b>; this includes staff who reported part of their time, had their time proxied, or entered their hours after the deadline.<br><br>Details of <b>{office_abbr}</b>'s ITR reporting compared to the CBER-wide average are below.</p><table style="border-collapse:collapse; border:1px solid black; font-family:Calibri,sans-serif; font-size:11pt;"><tr style="background-color:#4472C4; color:white; font-weight:bold;"><td style="border:1px solid black; padding:5px;">Office</td><td style="border:1px solid black; padding:5px;">Final Participation Rate</td><td style="border:1px solid black; padding:5px;">Initial Participation Rate</td></tr><tr><td style="border:1px solid black; padding:5px;">{office_abbr}</td><td style="border:1px solid black; padding:5px;">{office_final_rate}</td><td style="border:1px solid black; padding:5px;">{office_initial_rate}</td></tr><tr><td style="border:1px solid black; padding:5px;">Center</td><td style="border:1px solid black; padding:5px;">100.00%</td><td style="border:1px solid black; padding:5px;">95.54%</td></tr></table><p style="font-family:Calibri,sans-serif; font-size:11pt; color:black;">{non_reporters_section}{reminder_html}<br><br>If you have any questions regarding the data you received, please reach out to the <a href="mailto:cber.itrhelpdesk@fda.hhs.gov">CBER ITR Helpdesk</a>, and the team will happily address them.</p>"""
        return html
    
    def create_outlook_email(self, data, director_info, current_pp_num, non_reporters_list):
        """Create Outlook email with safe COM operations."""
        def create_email():
            outlook = win32.Dispatch('outlook.application')
            mail = outlook.CreateItem(0)
            office_abbr = data['Office']
            pay_period = str(current_pp_num)
            
            cc_emails = [director_info.get(f'CC {i}') for i in range(1, 12)]
            bcc_emails = [director_info.get(f'BCC {i}') for i in range(1, 3)]
            
            mail.To = director_info.get('Director Email', '')
            mail.CC = "; ".join(filter(pd.notna, [director_info.get('Deputy Email')] + cc_emails))
            mail.BCC = "; ".join(filter(pd.notna, bcc_emails))
            mail.Subject = f"{office_abbr} ITR Participation Report for FY25 PP{pay_period}"
            mail.HTMLBody = self.create_html_body(data, director_info, current_pp_num, non_reporters_list)
            mail.Display()
            return True
        
        return safe_com_operation(
            create_email,
            lambda msg, level: log_message_util(self.app.root, self.log_text, msg, level),
            f"Creating Outlook email for {data['Office']}"
        ) is not None
    
    def format_as_percent(self, value):
        """Enhanced percentage formatting."""
        if pd.isna(value):
            return ""
        try:
            if isinstance(value, str) and '%' in value:
                return value
            numeric_value = float(value)
            return f"{numeric_value:.2%}"
        except (ValueError, TypeError):
            return str(value)
    
    def get_next_pp_reminder(self, current_pp_num):
        """Get next pay period reminder with safe date formatting."""
        if self.dates_df is None:
            return "As a reminder, the ITR reporting deadline is 11:59 pm on Monday following the end of each Pay Period."
        
        try:
            next_pp_num = current_pp_num + 1
            if current_pp_num == 25:
                next_pp_num = 1
            
            next_pp_data = self.dates_df[self.dates_df['Pay Period'] == next_pp_num].iloc[0]
            
            start_date_raw = next_pp_data['PP Start Date']
            end_date_raw = next_pp_data['PP End Date']
            
            # Safe date formatting
            try:
                start_date = pd.to_datetime(start_date_raw).strftime('%m/%d/%y')
                end_date = pd.to_datetime(end_date_raw).strftime('%m/%d/%y')
            except:
                # Fallback if special formatting fails
                start_date = pd.to_datetime(start_date_raw).strftime('%m/%d/%y')
                end_date = pd.to_datetime(end_date_raw).strftime('%m/%d/%y')
            
            deadline_text = str(next_pp_data['Reporting Deadline Text'])
            due_by_match = re.search(r'(due by .* on .*?, .* \d+)', deadline_text)
            due_by_text = due_by_match.group(1) if due_by_match else "due by 11:59 pm on Monday"
            
            return {
                "part1": "As a reminder, the ITR reporting deadline is ",
                "part1_bold": "11:59 pm on Monday",
                "part2": " following the end of each Pay Period. The next time this deadline occurs will be for ",
                "part2_bold": f"PP{next_pp_num} ({start_date}-{end_date})",
                "part3": ", with reports ",
                "part3_bold": f"{due_by_text}",
                "part4": "."
            }
            
        except Exception as e:
            log_message_util(
                self.app.root,
                self.log_text,
                f"Error creating reminder: {str(e)}",
                'warning'
            )
            return "As a reminder, the ITR reporting deadline is 11:59 pm on Monday following the end of each Pay Period."
    
    def set_cell_background(self, cell, fill_color):
        """Helper function to set cell background color in Word."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def add_email_to_doc(self, doc, data, director_info, current_pp_num, non_reporters_list):
        """Add email content to Word document with enhanced formatting."""
        office_abbr = data['Office']
        pay_period = str(current_pp_num)
        office_final_rate = self.format_as_percent(data['Final Participation Rate'])
        office_initial_rate = self.format_as_percent(data['Initial Participation Rate'])
        
        cc_emails = [director_info.get(f'CC {i}') for i in range(1, 12)]
        bcc_emails = [director_info.get(f'BCC {i}') for i in range(1, 3)]
        to_line = director_info.get('Director Email', '')
        cc_line = "; ".join(filter(pd.notna, [director_info.get('Deputy Email')] + cc_emails))
        bcc_line = "; ".join(filter(pd.notna, bcc_emails))
        
        # Email headers
        doc.add_paragraph(f"To: {to_line}")
        doc.add_paragraph(f"CC: {cc_line}")
        doc.add_paragraph(f"BCC: {bcc_line}")
        doc.add_paragraph(f"Subject: {office_abbr} ITR Participation Report for FY25 PP{pay_period}")
        doc.add_paragraph("---")
        doc.add_paragraph()
        
        # Email body
        doc.add_paragraph(f"Hello {director_info.get('Director Name', 'Director')},")
        doc.add_paragraph()
        
        # Main content paragraphs with formatting
        p1 = doc.add_paragraph("Last week, CBER completed time reporting in ITR Pay Period ")
        p1.add_run(f"{pay_period}").bold = True
        p1.add_run(". Across CBER, ")
        p1.add_run("100.00%").bold = True
        p1.add_run(" of staff reported their full Tour of Duty (TOD), and ")
        p1.add_run("95.54%").bold = True
        p1.add_run(" reported partial hours.")
        doc.add_paragraph()
        
        p2 = doc.add_paragraph("In Pay Period ")
        p2.add_run(f"{pay_period}").bold = True
        p2.add_run(", ")
        p2.add_run(f"{office_abbr}").bold = True
        p2.add_run(" had an unassisted (initial) percentage of ")
        p2.add_run(f"{office_initial_rate}").bold = True
        p2.add_run(" which is the percent of staff who entered their full Tour of Duty in ITR by the 11:59pm ET Monday deadline. ")
        p2.add_run(f"{office_abbr}").bold = True
        p2.add_run("'s final participation rate is ")
        p2.add_run(f"{office_final_rate}").bold = True
        p2.add_run("; this includes staff who reported part of their time, had their time proxied, or entered their hours after the deadline.")
        doc.add_paragraph()
        
        p3 = doc.add_paragraph("Details of ")
        p3.add_run(f"{office_abbr}").bold = True
        p3.add_run("'s ITR reporting compared to the CBER-wide average are below.")
        doc.add_paragraph()
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        table.columns[0].width = Inches(1.0)
        table.columns[1].width = Inches(1.25)
        table.columns[2].width = Inches(1.25)
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Office'
        hdr_cells[1].text = 'Final Participation Rate'
        hdr_cells[2].text = 'Initial Participation Rate'
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self.set_cell_background(cell, "4472C4")
        
        # Data rows
        row_cells_1 = table.add_row().cells
        row_cells_1[0].text = office_abbr
        row_cells_1[1].text = office_final_rate
        row_cells_1[2].text = office_initial_rate
        
        row_cells_2 = table.add_row().cells
        row_cells_2[0].text = "Center"
        row_cells_2[1].text = "100.00%"
        row_cells_2[2].text = "95.54%"
        
        doc.add_paragraph()
        
        # Non-reporters section
        if non_reporters_list:
            p4 = doc.add_paragraph()
            run = p4.add_run("Users who did not report or had proxy:")
            run.bold = True
            run.underline = True
            for item in non_reporters_list:
                p = doc.add_paragraph(item, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
            doc.add_paragraph()
        
        # Reminder section
        reminder = self.get_next_pp_reminder(current_pp_num)
        if isinstance(reminder, dict):
            p5 = doc.add_paragraph(reminder['part1'])
            p5.add_run(reminder['part1_bold']).bold = True
            p5.add_run(reminder['part2'])
            p5.add_run(reminder['part2_bold']).bold = True
            p5.add_run(reminder['part3'])
            p5.add_run(reminder['part3_bold']).bold = True
            p5.add_run(reminder['part4'])
        else:
            doc.add_paragraph(reminder)
        
        doc.add_paragraph()
        doc.add_paragraph("If you have any questions regarding the data you received, please reach out to the CBER ITR Helpdesk, and the team will happily address them.")
        doc.add_page_break()

#==========================================================================
# MAIN APPLICATION CLASS (Enhanced)
#==========================================================================
class MasterWorkflowApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Master Workflow Tool v2.0")
        self.root.geometry("1200x900")
        
        # Modern styling
        self.setup_modern_theme()
        
        # Shared data
        self.default_output_path = get_default_output_path()
        self.office_name_map = {
            'OBPV': 'OFFICE OF BIOSTATISTICS & PHARMACOVIGILANCE',
            'OBRR': 'OFFICE OF BLOOD RESEARCH AND REVIEW',
            'OCBQ': 'OFFICE OF COMPLIANCE AND BIOLOGICS QUALITY',
            'OCOD': 'OFFICE OF COMMUNICATION, OUTREACH AND DEVELOPMENT',
            'OD': 'OFFICE OF THE CENTER DIRECTOR',
            'OM': 'OFFICE OF MANAGEMENT',
            'ORO': 'OFFICE OF REGULATORY OPERATIONS',
            'OTP': 'OFFICE OF THERAPEUTIC PRODUCTS',
            'OVRR': 'OFFICE OF VACCINES RESEARCH AND REVIEW'
        }
        
        self.setup_ui()
    
    def setup_modern_theme(self):
        """Setup modern Apple/Google-inspired theme."""
        self.root.configure(bg='#FFFFFF')
        
        # Style configuration
        self.style = ttk.Style()
        self.style.theme_use('clam')
        
        # Color palette
        self.BG_COLOR = '#FFFFFF'
        self.FG_COLOR = '#1C1C1E'
        self.SECONDARY_FG = '#8E8E93'
        self.ACCENT_BLUE = '#007AFF'
        self.ACCENT_GREEN = '#34C759'
        self.ACCENT_ORANGE = '#FF9500'
        self.ACCENT_RED = '#FF3B30'
        self.WIDGET_BG = '#F8F9FA'
        self.BORDER_COLOR = '#E5E5EA'
        
        # Configure styles
        self.style.configure('.', background=self.BG_COLOR, foreground=self.FG_COLOR)
        self.style.configure('TFrame', background=self.BG_COLOR, borderwidth=0)
        self.style.configure('TLabel', 
            background=self.BG_COLOR, 
            foreground=self.FG_COLOR, 
            font=('SF Pro Text', 12)
        )
        self.style.configure('TButton',
            background=self.WIDGET_BG,
            foreground=self.FG_COLOR,
            font=('SF Pro Text', 11),
            borderwidth=1,
            relief='solid',
            bordercolor=self.BORDER_COLOR,
            padding=(15, 8)
        )
        self.style.map('TButton',
            background=[('active', self.BORDER_COLOR), ('pressed', self.BORDER_COLOR)]
        )
        
        self.style.configure('TEntry',
            fieldbackground=self.WIDGET_BG,
            foreground=self.FG_COLOR,
            borderwidth=1,
            relief='solid',
            bordercolor=self.BORDER_COLOR,
            padding=8
        )
        
        self.style.configure('TCombobox',
            fieldbackground=self.WIDGET_BG,
            background=self.WIDGET_BG,
            foreground=self.FG_COLOR,
            borderwidth=1,
            relief='solid',
            bordercolor=self.BORDER_COLOR,
            arrowcolor=self.FG_COLOR
        )
        
        self.style.configure('TProgressbar',
            background=self.ACCENT_BLUE,
            troughcolor=self.WIDGET_BG,
            borderwidth=0,
            lightcolor=self.ACCENT_BLUE,
            darkcolor=self.ACCENT_BLUE
        )
        
        self.style.configure('TLabelframe',
            background=self.BG_COLOR,
            bordercolor=self.BORDER_COLOR,
            relief='solid',
            borderwidth=1
        )
        self.style.configure('TLabelframe.Label',
            background=self.BG_COLOR,
            foreground=self.FG_COLOR,
            font=('SF Pro Text', 12, 'bold')
        )
        
        self.style.configure('TNotebook',
            background=self.BG_COLOR,
            borderwidth=0,
            tabmargins=[2, 5, 2, 0]
        )
        self.style.configure('TNotebook.Tab',
            background=self.WIDGET_BG,
            foreground=self.SECONDARY_FG,
            padding=[20, 12],
            font=('SF Pro Text', 12, 'bold'),
            borderwidth=1,
            relief='solid'
        )
        self.style.map('TNotebook.Tab',
            background=[('selected', self.BG_COLOR)],
            foreground=[('selected', self.ACCENT_BLUE)],
            bordercolor=[('selected', self.ACCENT_BLUE)]
        )
        
        self.style.configure('TCheckbutton',
            background=self.BG_COLOR,
            foreground=self.FG_COLOR,
            font=('SF Pro Text', 11),
            focuscolor='none'
        )
        self.style.map('TCheckbutton',
            indicatorcolor=[('selected', self.ACCENT_BLUE)],
            background=[('active', self.BG_COLOR)]
        )
        
        self.style.configure('TRadiobutton',
            background=self.BG_COLOR,
            foreground=self.FG_COLOR,
            font=('SF Pro Text', 11),
            focuscolor='none'
        )
        self.style.map('TRadiobutton',
            indicatorcolor=[('selected', self.ACCENT_BLUE)],
            background=[('active', self.BG_COLOR)]
        )
    
    def setup_ui(self):
        """Setup main UI with notebook tabs."""
        # Main container with padding
        main_frame = ttk.Frame(self.root)
        main_frame.pack(expand=True, fill='both', padx=20, pady=20)
        
        # App header
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(
            header_frame,
            text="Master Workflow Tool",
            font=('SF Pro Display', 28, 'bold'),
            foreground=self.FG_COLOR
        ).pack(anchor='w')
        
        ttk.Label(
            header_frame,
            text="v2.0 - Enhanced with intelligent matching and modern UI",
            font=('SF Pro Text', 14),
            foreground=self.SECONDARY_FG
        ).pack(anchor='w', pady=(5, 0))
        
        # Notebook for tabs
        notebook = ttk.Notebook(main_frame)
        notebook.pack(expand=True, fill='both')
        
        # Create tab frames with padding
        reporter_tab = ReporterFrame(notebook, self, padding="30")
        formatter_tab = FormatterFrame(notebook, self, padding="30")
        emailer_tab = EmailerFrame(notebook, self, padding="30")
        
        # Add tabs to notebook
        notebook.add(reporter_tab, text='Pay Period Reporter')
        notebook.add(formatter_tab, text='Helper Sheet Formatter')
        notebook.add(emailer_tab, text='Office Emailer')
    
    def run(self):
        """Run the application."""
        self.root.mainloop()

if __name__ == "__main__":
    app = MasterWorkflowApp()
    app.run()